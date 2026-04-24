"""
Convert docs/phase3_report.md to docs/phase3_report.docx.

Academic black-and-white format: Times New Roman body, bold black headings,
plain bordered tables, no color accents.
"""
from __future__ import annotations
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD_PATH = os.path.join(ROOT, "docs", "phase3_report.md")
OUT_PATH = os.path.join(ROOT, "docs", "phase3_report.docx")

BLACK = RGBColor(0x00, 0x00, 0x00)
BODY_FONT = "Times New Roman"
CODE_FONT = "Courier New"
BODY_SIZE = Pt(11)
CODE_SIZE = Pt(9.5)


def add_hr(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    p_pr.append(pbdr)


def set_cell_borders(cell, color="000000", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def style_run(run, *, font=BODY_FONT, size=BODY_SIZE, bold=False, italic=False):
    run.font.name = font
    run.font.size = size
    run.font.color.rgb = BLACK
    run.bold = bold
    run.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)


def add_inline_runs(paragraph, text, base_size=BODY_SIZE):
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            style_run(r, size=base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            style_run(r, font=CODE_FONT, size=Pt(base_size.pt - 1))
        else:
            r = paragraph.add_run(part)
            style_run(r, size=base_size)


def is_table_row(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")

def is_table_separator(line: str) -> bool:
    s = line.strip().strip("|").strip()
    if not s:
        return False
    cells = [c.strip() for c in s.split("|")]
    return all(re.fullmatch(r":?-+:?", c) for c in cells)

def split_row(line: str):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def add_cover_page(doc):
    # Large vertical whitespace before title
    for _ in range(6):
        doc.add_paragraph()

    # Project title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("PuzzleForge")
    style_run(r, size=Pt(32), bold=True)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("An Agentic Puzzle Game Design and Generation System")
    style_run(r, size=Pt(16), italic=True)

    # Report label
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Phase 3 Final Report")
    style_run(r, size=Pt(14), bold=True)

    for _ in range(4):
        doc.add_paragraph()

    # Metadata block (centered, no colons or labels on separate lines)
    meta_lines = [
        ("Course", "94-815 Agentic Systems Studio"),
        ("Instructor", "Prof. Raj Sharman"),
        ("Team", "Yixiao Li, Kaizhen Tan, Hanzhe Hong"),
        ("Track", "A -- Technical Build"),
        ("College", "H. John Heinz III College"),
        ("University", "Carnegie Mellon University"),
        ("Submission Date", "April 24, 2026"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        rl = p.add_run(f"{label}:  ")
        style_run(rl, size=Pt(12), bold=True)
        rv = p.add_run(value)
        style_run(rv, size=Pt(12))

    add_page_break(doc)


def convert(md_text: str) -> Document:
    doc = Document()

    # Base body style
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.font.color.rgb = BLACK

    # Page margins: 1 inch all around (academic standard)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Cover page
    add_cover_page(doc)

    lines = md_text.splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.rstrip()

        if not stripped.strip():
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].rstrip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run("\n".join(code_lines))
            style_run(run, font=CODE_FONT, size=CODE_SIZE)
            continue

        # Horizontal rule: skip entirely (no visual divider between sections)
        if re.fullmatch(r"-{3,}", stripped.strip()):
            i += 1
            continue

        # Pipe table
        if is_table_row(stripped) and (i + 1) < n and is_table_separator(lines[i + 1]):
            header = split_row(stripped)
            i += 2
            rows = []
            while i < n and is_table_row(lines[i].rstrip()):
                rows.append(split_row(lines[i].rstrip()))
                i += 1
            tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            for c, text in enumerate(header):
                hdr[c].text = ""
                p = hdr[c].paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run(text)
                style_run(r, size=Pt(10), bold=True)
                set_cell_borders(hdr[c])
            for ri, row in enumerate(rows):
                cells = tbl.rows[1 + ri].cells
                for c, text in enumerate(row):
                    cells[c].text = ""
                    p = cells[c].paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    add_inline_runs(p, text, base_size=Pt(10))
                    set_cell_borders(cells[c])
            continue

        # Headings
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True
            if level == 1:
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(6)
                size = Pt(16)
            elif level == 2:
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(4)
                size = Pt(13)
            elif level == 3:
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(3)
                size = Pt(12)
            else:
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                size = Pt(11)
            r = p.add_run(text)
            style_run(r, size=size, bold=True)
            i += 1
            continue

        # Ordered list
        m = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m:
            text = m.group(3)
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, text)
            i += 1
            continue

        # Unordered list
        m = re.match(r"^(\s*)[-*+]\s+(.*)$", line)
        if m:
            text = m.group(2)
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, text)
            i += 1
            continue

        # Regular paragraph
        para_lines = [stripped]
        j = i + 1
        while j < n:
            nxt = lines[j].rstrip()
            if not nxt.strip():
                break
            if re.match(r"^#{1,4}\s", nxt):
                break
            if re.match(r"^\s*(\d+)\.\s", nxt) or re.match(r"^\s*[-*+]\s", nxt):
                break
            if nxt.startswith("```"):
                break
            if is_table_row(nxt):
                break
            if re.fullmatch(r"-{3,}", nxt.strip()):
                break
            para_lines.append(nxt)
            j += 1
        text = " ".join(para_lines)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        add_inline_runs(p, text)
        i = j

    return doc


def main() -> None:
    if not os.path.exists(MD_PATH):
        raise SystemExit(
            f"{MD_PATH} not found. Cannot rebuild .docx without the markdown source."
        )
    with open(MD_PATH, "r", encoding="utf-8") as f:
        md = f.read()
    doc = convert(md)
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
